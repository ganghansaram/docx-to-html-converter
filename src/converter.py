#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX to HTML Converter - 변환 로직
"""

import json
import os
import re
import hashlib
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx 패키지가 필요합니다. pip install python-docx 를 실행하세요.")

from utils import (
    get_logger, ensure_dir, get_image_dir, escape_html,
    convert_smart_quotes, ConversionResult, sanitize_filename
)


class DocxConverter:
    """Word 문서를 HTML로 변환하는 클래스"""

    def __init__(self, config_path=None):
        """
        Args:
            config_path: 설정 파일 경로 (기본값: ../config.json)
        """
        self.config = self._load_config(config_path)
        self.logger = get_logger()

    def _load_config(self, config_path):
        """설정 파일 로드"""
        if config_path is None:
            base_dir = Path(__file__).parent.parent
            config_path = base_dir / "config.json"

        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            return self._default_config()

    def _default_config(self):
        """기본 설정 반환"""
        return {
            "style_mapping": {
                "by_style": {
                    "제목 1": "h1", "제목 2": "h2", "제목 3": "h3",
                    "Heading 1": "h1", "Heading 2": "h2", "Heading 3": "h3",
                    "Title": "h1", "Normal": "p"
                },
                "by_font_size": {
                    "24": "h1", "18": "h2", "14": "h3", "default": "p"
                },
                "priority": "style_first"
            },
            "text_formatting": {
                "bold": "strong", "italic": "em", "underline": "u",
                "strikethrough": "del", "subscript": "sub", "superscript": "sup"
            },
            "special_blocks": {
                "note": ["NOTE", "참고", "비고"],
                "caution": ["CAUTION", "주의"],
                "warning": ["WARNING", "경고"]
            },
            "options": {
                "remove_headers_footers": True,
                "extract_images": True,
                "remove_empty_paragraphs": True,
                "convert_smart_quotes": True
            },
            "output": {
                "fragment_only": True,
                "encoding": "utf-8",
                "indent": True
            }
        }

    def convert(self, input_path, output_path=None, options=None):
        """
        Word 문서를 HTML로 변환

        Args:
            input_path: 입력 .docx 파일 경로
            output_path: 출력 .html 파일 경로 (선택)
            options: 변환 옵션 오버라이드 (선택)

        Returns:
            ConversionResult: 변환 결과 정보
        """
        result = ConversionResult(input_path)
        input_path = Path(input_path)

        # 입력 파일 검증
        if not input_path.exists():
            result.error_message = f"파일을 찾을 수 없습니다: {input_path}"
            return result

        if input_path.suffix.lower() != '.docx':
            result.error_message = f"지원하지 않는 파일 형식입니다: {input_path.suffix}"
            return result

        # 출력 경로 설정
        if output_path is None:
            output_path = input_path.with_suffix('.html')
        output_path = Path(output_path)
        result.output_path = output_path

        # 옵션 병합
        merged_options = {**self.config.get('options', {})}
        if options:
            merged_options.update(options)

        try:
            # Document 로드
            doc = Document(str(input_path))
            self.logger.info(f"문서 로드 완료: {input_path}")

            # 이미지 추출
            image_map = {}
            if merged_options.get('extract_images', True):
                image_dir = get_image_dir(output_path)
                image_map = self._process_images(doc, image_dir, output_path)
                result.stats['images'] = len(image_map)

            # HTML 생성
            html_parts = []
            first_heading_found = False

            for element in self._iter_block_items(doc):
                if hasattr(element, 'style') and hasattr(element, 'text'):
                    # 문단 처리
                    html, tag_type = self._process_paragraph(element, image_map, merged_options)

                    if html:
                        # h1 시작 여부 확인
                        if tag_type == 'h1':
                            first_heading_found = True
                        elif not first_heading_found and tag_type != 'h1':
                            result.add_warning("문서가 h1으로 시작하지 않습니다.")

                        html_parts.append(html)
                        result.stats['paragraphs'] += 1

                        if re.match(r'^h[1-6]$', tag_type):
                            result.stats['headings'][tag_type] = result.stats['headings'].get(tag_type, 0) + 1

                elif hasattr(element, 'rows'):
                    # 표 처리
                    table_html = self._process_table(element)
                    if table_html:
                        html_parts.append(table_html)
                        result.stats['tables'] += 1

            # HTML 결합
            indent = self.config.get('output', {}).get('indent', True)
            if indent:
                html_content = '\n\n'.join(html_parts)
            else:
                html_content = ''.join(html_parts)

            # h1으로 시작하지 않는 경우 경고
            if not first_heading_found:
                result.add_warning("문서에 h1 제목이 없습니다.")

            # 파일 저장
            ensure_dir(output_path.parent)
            encoding = self.config.get('output', {}).get('encoding', 'utf-8')

            with open(output_path, 'w', encoding=encoding) as f:
                f.write(html_content)

            self.logger.info(f"변환 완료: {output_path}")
            result.success = True

        except Exception as e:
            result.error_message = str(e)
            self.logger.error(f"변환 실패: {input_path} - {e}")

        return result

    def _iter_block_items(self, doc):
        """
        문서의 블록 요소(문단, 표)를 순서대로 반복

        Args:
            doc: Document 객체

        Yields:
            Paragraph 또는 Table 객체
        """
        for element in doc.element.body:
            if element.tag.endswith('p'):
                yield doc.paragraphs[
                    list(doc.element.body).index(element) -
                    len([e for e in list(doc.element.body)[:list(doc.element.body).index(element)]
                         if not e.tag.endswith('p')])
                ] if False else None  # placeholder
            elif element.tag.endswith('tbl'):
                pass

        # 단순화된 구현: 문단과 표를 순서대로 처리
        para_idx = 0
        table_idx = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                if para_idx < len(doc.paragraphs):
                    yield doc.paragraphs[para_idx]
                    para_idx += 1
            elif tag == 'tbl':
                if table_idx < len(doc.tables):
                    yield doc.tables[table_idx]
                    table_idx += 1

    def _detect_heading_level(self, paragraph):
        """
        문단의 제목 레벨 감지

        Args:
            paragraph: Paragraph 객체

        Returns:
            str: HTML 태그 (h1, h2, h3, p)
        """
        style_mapping = self.config.get('style_mapping', {})
        priority = style_mapping.get('priority', 'style_first')

        tag_from_style = None
        tag_from_font = None

        # 스타일 기반 감지
        by_style = style_mapping.get('by_style', {})
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name
            tag_from_style = by_style.get(style_name)

            # 매핑에 없으면 스타일명에서 숫자 자동 추출 (Heading 4, 제목 5 등)
            if not tag_from_style:
                match = re.match(r'^(?:Heading|제목)\s*(\d+)$', style_name, re.IGNORECASE)
                if match:
                    level = min(int(match.group(1)), 6)
                    tag_from_style = f'h{level}'

        # 폰트 크기 기반 감지
        by_font_size = style_mapping.get('by_font_size', {})
        font_size = self._get_paragraph_font_size(paragraph)
        if font_size:
            font_size_str = str(int(font_size))
            tag_from_font = by_font_size.get(font_size_str)

        # 우선순위에 따라 결정
        if priority == 'style_first':
            return tag_from_style or tag_from_font or by_font_size.get('default', 'p')
        else:
            return tag_from_font or tag_from_style or by_font_size.get('default', 'p')

    def _get_paragraph_font_size(self, paragraph):
        """
        문단의 폰트 크기 추출 (포인트 단위)

        Args:
            paragraph: Paragraph 객체

        Returns:
            float or None: 폰트 크기
        """
        # 첫 번째 run의 폰트 크기 확인
        for run in paragraph.runs:
            if run.font.size:
                return run.font.size.pt

        # 스타일에서 폰트 크기 확인
        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            return paragraph.style.font.size.pt

        return None

    def _process_paragraph(self, paragraph, image_map, options):
        """
        문단을 HTML로 변환

        Args:
            paragraph: Paragraph 객체
            image_map: 이미지 매핑 딕셔너리
            options: 변환 옵션

        Returns:
            tuple: (HTML 문자열, 태그 타입)
        """
        text = paragraph.text.strip()

        # 빈 문단 처리
        if not text and not self._has_images(paragraph):
            if options.get('remove_empty_paragraphs', True):
                return None, None

        # 특수 블록 감지
        special_block = self._detect_special_block(text)
        if special_block:
            block_type, content = special_block
            return f'<div class="{block_type}">{escape_html(content)}</div>', 'special'

        # 제목 레벨 감지
        tag = self._detect_heading_level(paragraph)

        # 정렬 속성 확인
        align_attr = ''
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_attr = ' style="text-align: center;"'
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_attr = ' style="text-align: right;"'

        # 인라인 서식 처리
        inner_html = self._process_runs(paragraph.runs, image_map, options)

        if not inner_html.strip():
            # 이미지만 있는 경우
            images_html = self._extract_inline_images(paragraph, image_map, align_attr)
            if images_html:
                return images_html, 'image'
            return None, None

        return f'<{tag}{align_attr}>{inner_html}</{tag}>', tag

    def _process_runs(self, runs, image_map, options):
        """
        Run 요소들을 HTML로 변환 (인라인 서식 처리)

        Args:
            runs: Run 객체 리스트
            image_map: 이미지 매핑
            options: 변환 옵션

        Returns:
            str: HTML 문자열
        """
        text_formatting = self.config.get('text_formatting', {})
        parts = []

        for run in runs:
            text = run.text

            if not text:
                # 인라인 이미지 확인
                inline_shape = self._get_run_image(run)
                if inline_shape and inline_shape in image_map:
                    img_path = image_map[inline_shape]
                    parts.append(f'<img src="{img_path}" alt="">')
                continue

            # 스마트 인용부호 변환
            if options.get('convert_smart_quotes', True):
                text = convert_smart_quotes(text)

            # HTML 이스케이프
            text = escape_html(text)

            # 인라인 서식 적용
            if run.bold:
                tag = text_formatting.get('bold', 'strong')
                text = f'<{tag}>{text}</{tag}>'

            if run.italic:
                tag = text_formatting.get('italic', 'em')
                text = f'<{tag}>{text}</{tag}>'

            if run.underline:
                tag = text_formatting.get('underline', 'u')
                text = f'<{tag}>{text}</{tag}>'

            if run.font.strike:
                tag = text_formatting.get('strikethrough', 'del')
                text = f'<{tag}>{text}</{tag}>'

            if run.font.subscript:
                tag = text_formatting.get('subscript', 'sub')
                text = f'<{tag}>{text}</{tag}>'

            if run.font.superscript:
                tag = text_formatting.get('superscript', 'sup')
                text = f'<{tag}>{text}</{tag}>'

            parts.append(text)

        return ''.join(parts)

    def _process_table(self, table):
        """
        표를 HTML table로 변환

        Args:
            table: Table 객체

        Returns:
            str: HTML table 문자열
        """
        rows_html = []

        for i, row in enumerate(table.rows):
            cells_html = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = escape_html(cell_text)

                # 첫 번째 행은 헤더로 처리
                if i == 0:
                    cells_html.append(f'<th>{cell_text}</th>')
                else:
                    cells_html.append(f'<td>{cell_text}</td>')

            row_tag = '<tr>' + ''.join(cells_html) + '</tr>'
            rows_html.append(row_tag)

        # thead/tbody 분리
        if len(rows_html) > 1:
            thead = f'<thead>{rows_html[0]}</thead>'
            tbody = '<tbody>' + ''.join(rows_html[1:]) + '</tbody>'
            return f'<table>{thead}{tbody}</table>'
        elif rows_html:
            return f'<table><tbody>{"".join(rows_html)}</tbody></table>'
        else:
            return ''

    def _process_images(self, doc, image_dir, output_path):
        """
        문서의 이미지를 추출하고 저장

        Args:
            doc: Document 객체
            image_dir: 이미지 저장 디렉토리
            output_path: HTML 출력 경로 (상대 경로 계산용)

        Returns:
            dict: {rId: 상대경로} 매핑
        """
        image_map = {}

        try:
            # 이미지 디렉토리 생성
            ensure_dir(image_dir)

            # 문서의 모든 관계(relationship)에서 이미지 추출
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        content_type = rel.target_part.content_type

                        # 확장자 결정
                        ext_map = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/gif': '.gif',
                            'image/bmp': '.bmp',
                            'image/tiff': '.tiff',
                            'image/x-emf': '.emf',
                            'image/x-wmf': '.wmf',
                        }
                        ext = ext_map.get(content_type, '.png')

                        # 파일명 생성 (해시 기반)
                        hash_name = hashlib.md5(image_data).hexdigest()[:12]
                        filename = f"image_{hash_name}{ext}"
                        image_path = image_dir / filename

                        # 이미지 저장
                        with open(image_path, 'wb') as f:
                            f.write(image_data)

                        # 상대 경로 계산
                        rel_path = os.path.relpath(image_path, output_path.parent)
                        rel_path = rel_path.replace('\\', '/')

                        image_map[rel_id] = rel_path
                        self.logger.debug(f"이미지 추출: {filename}")

                    except Exception as e:
                        self.logger.warning(f"이미지 추출 실패 ({rel_id}): {e}")

        except Exception as e:
            self.logger.error(f"이미지 처리 중 오류: {e}")

        return image_map

    def _has_images(self, paragraph):
        """문단에 이미지가 있는지 확인"""
        # XML에서 drawing 요소 확인
        drawing_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        for elem in paragraph._element.iter():
            if 'drawing' in elem.tag or 'pict' in elem.tag:
                return True
        return False

    def _get_run_image(self, run):
        """Run에서 이미지 rId 추출"""
        try:
            for elem in run._element.iter():
                if 'blip' in elem.tag:
                    embed = elem.get(qn('r:embed'))
                    if embed:
                        return embed
        except:
            pass
        return None

    def _extract_inline_images(self, paragraph, image_map, align_attr=''):
        """문단에서 인라인 이미지 추출"""
        images = []

        try:
            for elem in paragraph._element.iter():
                if 'blip' in elem.tag:
                    embed = elem.get(qn('r:embed'))
                    if embed and embed in image_map:
                        img_path = image_map[embed]
                        images.append(f'<img src="{img_path}" alt="">')
        except:
            pass

        if images:
            return f'<p{align_attr}>' + ''.join(images) + '</p>'
        return None

    def _detect_special_block(self, text):
        """
        특수 블록 (NOTE, WARNING, CAUTION) 감지

        Args:
            text: 문단 텍스트

        Returns:
            tuple: (블록 타입, 내용) 또는 None
        """
        special_blocks = self.config.get('special_blocks', {})

        for block_type, keywords in special_blocks.items():
            for keyword in keywords:
                # "NOTE: 내용" 또는 "NOTE - 내용" 패턴
                patterns = [
                    rf'^{re.escape(keyword)}\s*[:：\-]\s*(.+)$',
                    rf'^【{re.escape(keyword)}】\s*(.+)$',
                    rf'^\[{re.escape(keyword)}\]\s*(.+)$',
                ]

                for pattern in patterns:
                    match = re.match(pattern, text, re.IGNORECASE)
                    if match:
                        return (block_type, match.group(1).strip())

        return None

    def analyze(self, input_path):
        """
        문서 구조 분석 (미리보기용)

        Args:
            input_path: 입력 .docx 파일 경로

        Returns:
            dict: 문서 구조 정보
        """
        input_path = Path(input_path)

        if not input_path.exists():
            return {'error': f"파일을 찾을 수 없습니다: {input_path}"}

        try:
            doc = Document(str(input_path))

            # 기본 정보
            analysis = {
                'filename': input_path.name,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'headings': {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0, 'h5': 0, 'h6': 0},
                'images': 0,
                'styles_used': set(),
                'font_sizes_used': set(),
                'starts_with_h1': False,
                'warnings': []
            }

            # 이미지 수 계산
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    analysis['images'] += 1

            # 문단 분석
            first_content_found = False
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 스타일 수집
                if para.style and para.style.name:
                    analysis['styles_used'].add(para.style.name)

                # 폰트 크기 수집
                font_size = self._get_paragraph_font_size(para)
                if font_size:
                    analysis['font_sizes_used'].add(font_size)

                # 제목 레벨 분석
                tag = self._detect_heading_level(para)
                if re.match(r'^h[1-6]$', tag):
                    analysis['headings'][tag] = analysis['headings'].get(tag, 0) + 1

                    if not first_content_found:
                        if tag == 'h1':
                            analysis['starts_with_h1'] = True
                        first_content_found = True
                elif not first_content_found:
                    first_content_found = True

            # 경고 생성
            if not analysis['starts_with_h1']:
                analysis['warnings'].append("문서가 h1으로 시작하지 않습니다.")

            if analysis['headings']['h1'] == 0:
                analysis['warnings'].append("h1 제목이 없습니다.")

            if analysis['headings']['h1'] > 1:
                analysis['warnings'].append(f"h1 제목이 {analysis['headings']['h1']}개 있습니다. (권장: 1개)")

            # set을 list로 변환 (JSON 직렬화용)
            analysis['styles_used'] = sorted(list(analysis['styles_used']))
            analysis['font_sizes_used'] = sorted(list(analysis['font_sizes_used']))

            return analysis

        except Exception as e:
            return {'error': str(e)}


# 테스트용
if __name__ == "__main__":
    converter = DocxConverter()
    print("Config loaded:", json.dumps(converter.config, indent=2, ensure_ascii=False))
