#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""샘플 DOCX 문서 생성"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_sample_document():
    """테스트용 샘플 문서 생성"""
    doc = Document()

    # 제목 (h1)
    title = doc.add_heading('KF-21 보라매 개발 개요', level=1)

    # 본문
    doc.add_paragraph(
        'KF-21 보라매는 대한민국이 개발한 4.5세대 전투기로, '
        '한국항공우주산업(KAI)이 주관하여 개발하였습니다.'
    )

    # 소제목 (h2)
    doc.add_heading('개발 배경', level=2)

    para = doc.add_paragraph()
    para.add_run('KF-21 개발 사업은 ').bold = False
    para.add_run('2015년').bold = True
    para.add_run('에 본격적으로 시작되었습니다. ')
    run = para.add_run('차세대 전투기')
    run.italic = True
    para.add_run(' 확보를 위한 국가적 프로젝트입니다.')

    # 소소제목 (h3)
    doc.add_heading('주요 특징', level=3)

    # 서식이 있는 문단
    para2 = doc.add_paragraph()
    para2.add_run('스텔스 성능').underline = True
    para2.add_run('과 ')
    run2 = para2.add_run('초음속 비행')
    run2.bold = True
    run2.italic = True
    para2.add_run(' 능력을 갖추고 있습니다.')

    # 표
    doc.add_heading('제원 정보', level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # 헤더
    table.rows[0].cells[0].text = '항목'
    table.rows[0].cells[1].text = '제원'

    # 데이터
    table.rows[1].cells[0].text = '전장'
    table.rows[1].cells[1].text = '16.9m'
    table.rows[2].cells[0].text = '전폭'
    table.rows[2].cells[1].text = '11.2m'
    table.rows[3].cells[0].text = '최대속도'
    table.rows[3].cells[1].text = '마하 1.81'

    # 특수 블록
    doc.add_paragraph()
    doc.add_paragraph('NOTE: 본 문서는 테스트용 샘플입니다.')
    doc.add_paragraph('WARNING: 실제 기밀 정보가 포함되어 있지 않습니다.')

    # 또 다른 섹션
    doc.add_heading('개발 일정', level=2)

    doc.add_paragraph('시제기 출고: 2021년 4월')
    doc.add_paragraph('초도비행: 2022년 7월')
    doc.add_paragraph('양산 예정: 2026년')

    # 저장
    output_dir = Path(__file__).parent / 'samples'
    output_dir.mkdir(exist_ok=True)

    output_path = output_dir / 'sample_document.docx'
    doc.save(str(output_path))
    print(f'샘플 문서 생성 완료: {output_path}')

    return output_path

if __name__ == '__main__':
    create_sample_document()
